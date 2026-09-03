"""
DocumentProcessor — Spec Section 3/4/5 ka entry point (processing/)

File aayi → sahi processor chuno → text + citation-ready chunks do.

Har citation-ready chunk capture/transformation integrity metadata bhi carry kar
sakta hai. Native text capture aur OCR capture alag rehte hain; OCR confidence
ko source quality ya truth score mein merge nahi kiya jaata.
"""
from __future__ import annotations

import os
import re
from typing import Dict, List

from ..extraction_integrity import native_text_integrity
from .ocr_processor import OCRProcessor
from .pdf_processor import PDFProcessor
from .transcript_processor import TranscriptProcessor
from . import pdf_chunker

TEXT_EXTENSIONS = (".txt", ".md", ".markdown", ".text")
TRANSCRIPT_EXTENSIONS = (".vtt", ".srt")
_TAG_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.IGNORECASE | re.DOTALL)
_ANY_TAG_RE = re.compile(r"<[^>]+>")


class DocumentProcessor:
    name = "document"

    def __init__(self, ocr_lang: str = "eng", ocr_max_pages: int = 20):
        self.pdf = PDFProcessor()
        self.ocr = OCRProcessor(lang=ocr_lang)
        self.transcript = TranscriptProcessor()
        self.ocr_max_pages = ocr_max_pages

    def process(self, file_path: str, use_ocr: bool = True,
                question: str = "", size_bytes: int = 0,
                large: bool = False) -> Dict:
        base = {"ok": False, "error": "", "text": "", "chunks": [], "notes": [],
                "kind": "", "file": os.path.basename(file_path or "")}
        if not file_path or not os.path.exists(file_path):
            base["error"] = f"file nahi mili: {file_path}"
            return base

        extension = os.path.splitext(file_path)[1].lower()
        if extension == ".pdf":
            return self._process_pdf(file_path, use_ocr, base, question=question,
                                     size_bytes=size_bytes, large=large)
        if extension in TRANSCRIPT_EXTENSIONS:
            return self._process_transcript(file_path, base)
        if extension in TEXT_EXTENSIONS:
            return self._process_text(file_path, base)
        if extension == ".docx":
            return self._process_docx(file_path, base)
        if extension in (".html", ".htm"):
            return self._process_html(file_path, base)
        base["error"] = (f"'{extension}' format supported nahi hai. Supported: "
                         f"pdf, vtt, srt, txt, md, docx, html")
        return base

    @staticmethod
    def _native_chunk(locator: str, text: str, header: str,
                      *, engine: str = "native_text") -> Dict:
        return {
            "locator": locator,
            "text": text,
            "header": header,
            "extraction_integrity": native_text_integrity(engine=engine),
        }

    def _process_pdf(self, file_path: str, use_ocr: bool, base: Dict,
                     question: str = "", size_bytes: int = 0,
                     large: bool = False) -> Dict:
        base["kind"] = "pdf"
        if not size_bytes:
            try:
                size_bytes = os.path.getsize(file_path)
            except Exception:
                size_bytes = 0
        pages_hint = self.pdf.page_count(file_path)
        if large or pdf_chunker.is_large(size_bytes=size_bytes,
                                        page_count=pages_hint):
            return self._process_pdf_streaming(file_path, use_ocr, base,
                                               question=question,
                                               size_bytes=size_bytes)

        result = self.pdf.extract(file_path)
        base["notes"].append(self.pdf.coverage_note(result))
        if not result["extracted"]:
            base["error"] = result["error"] or "PDF se text nahi nikla"
            return base

        text_parts = [result["text"]] if result["text"] else []
        chunks = [
            self._native_chunk(
                f"p.{page['page']}", page["text"],
                f"[Source: {base['file']}, Page {page['page']}]",
                engine="pdf_native_text",
            )
            for page in result["pages"] if page["text"] and not page["scanned"]
        ]

        scanned = result.get("scanned_pages", [])
        if scanned and use_ocr:
            ocr_result = self.ocr.ocr_pdf_pages(file_path, scanned,
                                                max_pages=self.ocr_max_pages)
            base["notes"].append(self.ocr.note(ocr_result))
            if ocr_result.get("text"):
                text_parts.append(ocr_result["text"])
                chunks += [
                    {
                        "locator": f"p.{page['page']} (OCR)",
                        "text": page["text"],
                        "header": f"[Source: {base['file']}, Page {page['page']}]",
                        "extraction_integrity": dict(
                            page.get("extraction_integrity") or {}
                        ),
                    }
                    for page in ocr_result["pages"] if page.get("text")
                ]
        elif scanned:
            base["notes"].append(f"{len(scanned)} scanned pages skip hue (OCR off tha)")

        base["text"] = "\n\n".join(part for part in text_parts if part)
        base["chunks"] = chunks
        base["ok"] = bool(base["text"])
        if not base["ok"]:
            base["error"] = ("PDF mein padhne layak text nahi mila (poora document "
                             "scanned ho sakta hai aur OCR available nahi tha)")
        return base

    def _process_pdf_streaming(self, file_path: str, use_ocr: bool, base: Dict,
                               question: str = "", size_bytes: int = 0) -> Dict:
        base["kind"] = "pdf"
        base["streamed"] = True
        result = self.pdf.extract_relevant(file_path, question,
                                           size_bytes=size_bytes)
        base["notes"].append(self.pdf.coverage_note(result))
        base["selection"] = result.get("selection", {})
        base["page_count"] = result.get("page_count", 0)

        chunks = [
            self._native_chunk(
                c["locator"], c["text"], c.get("header", ""),
                engine="pdf_streamed_native_text",
            )
            for c in (result.get("chunks") or []) if c.get("text")
        ]
        text_parts = [result.get("text") or ""]

        scanned = result.get("scanned_pages") or []
        if scanned and use_ocr and not chunks:
            budget = min(self.ocr_max_pages, 5)
            ocr_result = self.ocr.ocr_pdf_pages(file_path, scanned[:budget],
                                                max_pages=budget)
            base["notes"].append(self.ocr.note(ocr_result))
            if ocr_result.get("text"):
                text_parts.append(ocr_result["text"])
                chunks += [
                    {
                        "locator": f"p.{page['page']} (OCR)",
                        "text": page["text"],
                        "header": f"[Source: {base['file']}, Page {page['page']}]",
                        "extraction_integrity": dict(
                            page.get("extraction_integrity") or {}
                        ),
                    }
                    for page in ocr_result["pages"] if page.get("text")
                ]
        elif scanned and not use_ocr:
            base["notes"].append(f"{len(scanned)} scanned pages skip hue (OCR off tha)")

        base["text"] = "\n\n".join(part for part in text_parts if part)
        base["chunks"] = chunks
        base["ok"] = bool(base["text"])
        if not base["ok"]:
            base["error"] = result.get("error") or (
                "badi PDF ke chune hue pages se bhi padhne layak text nahi mila")
        return base

    def _process_transcript(self, file_path: str, base: Dict) -> Dict:
        base["kind"] = "transcript"
        result = self.transcript.process_file(file_path)
        chunks = []
        for chunk in result.get("chunks", []) or []:
            item = dict(chunk)
            item.setdefault("extraction_integrity",
                            native_text_integrity(engine="transcript_native_text"))
            chunks.append(item)
        base["ok"] = result["ok"]
        base["error"] = result.get("error", "")
        base["text"] = result.get("text", "")
        base["chunks"] = chunks
        if result.get("duration_note"):
            base["notes"].append(result["duration_note"])
        return base

    def _read(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _chunk_plain(self, text: str, file_name: str,
                     chunk_chars: int = 1200) -> List[Dict]:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        chunks: List[Dict] = []
        buffer: List[str] = []
        size = 0
        index = 1

        def flush(idx: int):
            if not buffer:
                return
            chunks.append(self._native_chunk(
                f"part {idx}", "\n\n".join(buffer),
                f"[Source: {file_name}, Part {idx}]",
                engine="plain_native_text",
            ))

        for paragraph in paragraphs:
            if size + len(paragraph) > chunk_chars and buffer:
                flush(index)
                index += 1
                buffer, size = [], 0
            buffer.append(paragraph)
            size += len(paragraph)
        flush(index)
        return chunks

    def _process_text(self, file_path: str, base: Dict) -> Dict:
        base["kind"] = "text"
        try:
            text = self._read(file_path).strip()
        except Exception as exc:
            base["error"] = f"file padhi nahi gayi: {type(exc).__name__}: {exc}"
            return base
        base["text"] = text
        base["chunks"] = self._chunk_plain(text, base["file"])
        base["ok"] = bool(text)
        base["notes"].append(f"{len(text)} chars, {len(base['chunks'])} chunks")
        if not text:
            base["error"] = "file khaali hai"
        return base

    def _process_docx(self, file_path: str, base: Dict) -> Dict:
        base["kind"] = "docx"
        try:
            import docx
        except Exception as exc:
            base["error"] = f"python-docx nahi hai ({exc}) — pip install python-docx"
            return base
        try:
            document = docx.Document(file_path)
        except Exception as exc:
            base["error"] = f"docx khul nahi rahi: {type(exc).__name__}: {exc}"
            return base
        text = "\n\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
        base["text"] = text
        base["chunks"] = self._chunk_plain(text, base["file"])
        base["ok"] = bool(text)
        base["notes"].append(f"{len(document.paragraphs)} paragraphs")
        if not text:
            base["error"] = "docx mein text nahi mila"
        return base

    def _process_html(self, file_path: str, base: Dict) -> Dict:
        base["kind"] = "html"
        try:
            raw = self._read(file_path)
        except Exception as exc:
            base["error"] = f"file padhi nahi gayi: {type(exc).__name__}: {exc}"
            return base
        without_scripts = _TAG_RE.sub(" ", raw)
        text = _ANY_TAG_RE.sub(" ", without_scripts)
        text = re.sub(r"&nbsp;?", " ", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        base["text"] = text
        base["chunks"] = self._chunk_plain(text, base["file"])
        base["ok"] = bool(text)
        if not text:
            base["error"] = "HTML se text nahi nikla"
        return base
